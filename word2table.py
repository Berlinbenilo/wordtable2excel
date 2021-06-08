import xlsxwriter
import os
from docx import Document
from tkinter.filedialog import askopenfilename
import tkinter.messagebox
from tkinter import *

root = Tk()
root.geometry('200x200')
root.title('Wordtable to excel sheet')


def select():
    global book, doc
    filename = askopenfilename(filetypes=[('all files','*.*')])
    doc = Document(filename)
    base = os.path.basename(filename)
    name = os.path.splitext(base)[0]
    book = str(name)+'.xlsx'

def convert():
    with xlsxwriter.Workbook(book) as workbook:
        for i in range(len(doc.tables)):
            tb = doc.tables[i]

            ls =[]
            main = []
            rows = tb.rows
            # print(len(rows),len(rows[0].cells))
            cols = rows[0].cells


            for row in rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        ls.append(para.text)
                        if len(ls) == len(rows[0].cells):
                            main.append(ls)
                            ls = []

            worksheet = workbook.add_worksheet()

            for row_num, data in enumerate(main):
                worksheet.write_row(row_num,0,data)
        tkinter.messagebox.showinfo('Success','Table converted Successfully')

def Exit():
    root.destroy()

Button1 = Button(root,text='Select doc',width=20,height=2,command=select,fg='blue',bg='red')
Button1.pack()
Button1.place(x=20,y=40)

Button2 = Button(root,text='Convert',width=20,height=2,command=convert,fg='blue',bg='red')
Button2.pack()
Button2.place(x=20,y=90)

Button2 = Button(root,text='Exit',width=20,height=2,command=Exit,fg='blue',bg='red')
Button2.pack()
Button2.place(x=20,y=140)

root.mainloop()
